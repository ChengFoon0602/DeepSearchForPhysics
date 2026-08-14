import logging
from pathlib import Path

try:
    from typing import Annotated, Optional
except ImportError:
    from typing_extensions import Annotated, Optional

import re

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

from langchain_core.tools import tool
from api.monitor import monitor
from api.context import get_session_context
from utils.path_utils import resolve_path


def _add_inline_runs(paragraph, text: str):
    """把一段文本按 **粗体** 和 `代码` 拆成多个 run 追加到段落。

    简单解析：**...** → 加粗；`...` → 等宽。其余原样。
    公式 $...$ 保留为文本（与 PDF 相同的诚实限制）。
    """
    # 先处理 **bold**，再处理 `code`
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _md_to_docx(md_text: str, doc) -> None:
    """把 Markdown 内容写进 python-docx 文档（标题/列表/表格/代码块/段落）。"""
    lines = md_text.splitlines()
    i = 0
    n = len(lines)
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        if not table_buf:
            return
        # 第一行是表头，第二行是分隔（|---|）
        header = [c.strip() for c in table_buf[0].strip().strip("|").split("|")]
        rows = []
        for row in table_buf[2:]:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            rows.append(cells)
        tbl = doc.add_table(rows=1 + len(rows), cols=max(1, len(header)))
        tbl.style = "Table Grid"
        for j, cell in enumerate(header):
            tbl.rows[0].cells[j].text = cell
        for ri, row in enumerate(rows):
            for j in range(len(header)):
                tbl.rows[ri + 1].cells[j].text = row[j] if j < len(row) else ""
        table_buf.clear()

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                # 代码块结束
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
                table_buf = []  # 代码块里不是表格
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格行
        if stripped.startswith("|") and stripped.endswith("|"):
            if table_buf and stripped.replace("|", "").strip().startswith("-"):
                table_buf.append(stripped)  # 分隔行
            else:
                flush_table()
                table_buf = [stripped]
            i += 1
            continue

        flush_table()

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            _add_inline_runs(heading, text)
            i += 1
            continue

        # 列表项
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, m.group(1).strip())
            i += 1
            continue
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, m.group(1).strip())
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    flush_table()


@tool
def generate_docx(
        md_filename: Annotated[str, "源 Markdown 文档路径（包含 .md 后缀）"],
        docx_filename: Annotated[Optional[str], "输出的 Word 文档路径（可选，默认与源文件同名）"] = None
) -> str:
    """
    将 Markdown 文档转换为 Word (.docx) 文件（基于 python-docx）。
    支持标题/列表/表格/代码块/粗体/行内代码；公式 $...$ 保留为文本（与 PDF 同限制）。
    """
    monitor.report_tool("Word文档生成工具", {"md_filename": md_filename, "docx_filename": docx_filename})

    if docx is None:
        return "错误：缺少 python-docx 依赖，请安装: pip install python-docx"

    try:
        # 路径解析（同 pdf_tools 模式）
        session_dir = get_session_context()
        md_path = Path(md_filename).with_suffix(".md")
        md_abs_path = Path(resolve_path(str(md_path), session_dir))
        if not md_abs_path.exists():
            return f"错误：文件不存在 {md_abs_path}"

        if docx_filename:
            docx_path = Path(docx_filename).with_suffix(".docx")
            docx_abs_path = Path(resolve_path(str(docx_path), session_dir))
        else:
            docx_abs_path = md_abs_path.with_suffix(".docx")

        # 生成
        doc = docx.Document()
        md_text = md_abs_path.read_text(encoding="utf-8")
        _md_to_docx(md_text, doc)
        doc.save(str(docx_abs_path))

        if docx_abs_path.exists():
            return f"成功生成 Word 文档: {docx_abs_path}"
        return f"转换完成但未生成文件: {docx_abs_path}"
    except Exception as e:
        logging.error(f"Word文档生成失败: {e}", exc_info=True)
        return f"生成 Word 文档失败: {str(e)}"


if __name__ == "__main__":
    # 测试：用 reports/ 里一份 md 生成 docx
    import sys
    sys.path.insert(0, str(Path(__file__).parents[1]))
    import os
    os.chdir(Path(__file__).parents[1])
    from api.context import set_session_context
    set_session_context(str(Path(__file__).parents[1] / "reports"))
    print(generate_docx.invoke({"md_filename": "复摆周期公式推导与验证报告.md"}))
